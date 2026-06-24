import asyncio
import io
from typing import Any

import structlog

logger = structlog.get_logger()


async def parse_docx(file_bytes: bytes) -> str:
    def _extract() -> str:
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            blocks: list[str] = []

            # Body paragraphs (python-docx's doc.paragraphs excludes text inside
            # tables, so tables are gathered separately below without duplication).
            for para in doc.paragraphs:
                if para.text.strip():
                    blocks.append(para.text.strip())

            # Tables: render each row as tab-joined cells so structure survives.
            for table in doc.tables:
                rows = _table_rows(table)
                if rows:
                    blocks.append("\n".join(rows))

            # Headers and footers (per section, primary + first-page + even).
            for section in doc.sections:
                for hf in (section.header, section.footer):
                    for para in hf.paragraphs:
                        if para.text.strip():
                            blocks.append(para.text.strip())
                    for table in hf.tables:
                        rows = _table_rows(table)
                        if rows:
                            blocks.append("\n".join(rows))

            return "\n\n".join(blocks)
        except Exception as e:
            logger.error("docx_parse_error", error=str(e))
            return ""

    return await asyncio.to_thread(_extract)


def _table_rows(table: Any) -> list[str]:
    """Render a docx table as a list of tab-joined non-empty rows."""
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            rows.append("\t".join(cells))
    return rows
